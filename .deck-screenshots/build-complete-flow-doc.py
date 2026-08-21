from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = r"D:\Projects\AI AGENT\GIT_MCP_AGENT\gitarchitect\backend\GitArchitect_Complete_Application_Flow.docx"


COLORS = {
    "ink": RGBColor(7, 20, 43),
    "muted": RGBColor(82, 98, 122),
    "blue": RGBColor(46, 116, 181),
    "dark_blue": RGBColor(31, 77, 120),
    "border": "D5DEEB",
    "fill": "F2F6FB",
    "callout": "ECFBF3",
    "warn": "FFF7DF",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D5DEEB", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def style_table(table, widths=None, header=True):
    if widths:
        set_table_width(table, widths)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_border(cell)
            set_cell_margins(cell)
            if header and row_idx == 0:
                set_cell_shading(cell, COLORS["fill"])
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(3)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        run.font.color.rgb = COLORS["blue"] if level < 3 else COLORS["dark_blue"]
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    return p


def add_para(doc, text="", bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.add_run(item)


def add_callout(doc, title, body, fill="ECFBF3"):
    table = doc.add_table(rows=1, cols=1)
    style_table(table, [6.5], header=False)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = COLORS["ink"]
    p.add_run(" " + body)
    doc.add_paragraph()


def add_kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Area"
    table.cell(0, 1).text = "Details"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    style_table(table, [1.75, 4.75])
    doc.add_paragraph()


def add_route_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    headers = ["Page / Feature", "Frontend action", "Backend endpoint"]
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    style_table(table, [1.55, 2.2, 2.75])
    doc.add_paragraph()


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("GitArchitect Complete Application Flow")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = COLORS["ink"]

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("Frontend pages, backend APIs, session flow, GitHub integration, OpenAI agents, and presentation notes.")
    r.font.size = Pt(12)
    r.font.color.rgb = COLORS["muted"]

    add_callout(
        doc,
        "Executive summary.",
        "GitArchitect is an AI engineering assistant for GitHub repositories. The Angular frontend lets the user select a repository, ask questions, analyze repository quality, review pull requests, analyze issues, debug CI failures, and prepare GitHub write actions. The Express backend keeps the active session and selected repository in MongoDB, then calls OpenAI agents with GitHub MCP/API evidence so responses are repository-aware.",
    )

    add_heading(doc, "1. System Overview")
    add_kv_table(
        doc,
        [
            ("Frontend", "Angular application with pages for Dashboard, Repositories, AI Chat, Repository Intelligence, PR Review, Issue Analyzer, CI Debugger, and GitHub Actions."),
            ("Backend", "Node.js, Express 5, TypeScript API under /api/v1/agent and /api/v1/github."),
            ("State", "MongoDB stores agent sessions, selected repository context, and GitHub write approvals."),
            ("AI layer", "OpenAI Agents SDK runs specialist agents for chat, repository analysis, PR review, issue analysis, CI debugging, and GitHub write planning."),
            ("GitHub layer", "GitHub REST and MCP tools provide repository metadata, source files, pull requests, issues, checks, workflow runs, and logs."),
        ],
    )

    add_heading(doc, "2. High-Level Application Flow")
    add_numbered(
        doc,
        [
            "User opens the Angular app and the frontend creates or reuses an agent session.",
            "User connects GitHub or checks connection status.",
            "User selects a repository from the Repositories page.",
            "Frontend sends the selected repository to the backend with sessionId.",
            "Backend verifies repository access and stores repository context in MongoDB.",
            "All analysis pages send workflow-specific input plus sessionId.",
            "Backend resolves the selected repository from MongoDB before calling any agent.",
            "OpenAI agents inspect repository evidence using GitHub MCP/API tools.",
            "Backend returns structured JSON to the frontend.",
            "Frontend renders the response in cards: summary, risks, scores, findings, plans, files, limitations, and questions.",
        ],
    )

    add_callout(
        doc,
        "Important rule.",
        "sessionId and selected repository are mandatory for repository-specific workflows. Without them, owner/repo becomes undefined and GitHub calls can fail or UI cards can show blank values.",
        fill=COLORS["warn"],
    )

    add_heading(doc, "3. Main Pages and Responsibilities")
    add_route_table(
        doc,
        [
            ("Dashboard", "Shows current session, selected repository, and GitHub write mode.", "GET /api/v1/github/auth/status and session state from frontend storage/API."),
            ("Repositories", "Searches available GitHub repositories and selects one.", "GET /api/v1/github/repos; PUT /api/v1/agent/sessions/:sessionId/repository."),
            ("AI Chat", "Allows questions about the active repository.", "POST /api/v1/agent/chat."),
            ("Repository Intelligence", "Runs architecture, security, testing, and stack analysis.", "POST /api/v1/agent/sessions/:sessionId/repository-analysis."),
            ("PR Review", "Reviews a pull request and shows score, recommendation, findings, checks, and limitations.", "POST /api/v1/agent/sessions/:sessionId/pull-requests/:pullNumber/review."),
            ("Issue Analyzer", "Converts a GitHub issue into requirements, plan, affected areas, risks, questions, and files.", "POST /api/v1/agent/sessions/:sessionId/issues/:issueNumber/analyze."),
            ("CI Debugger", "Investigates failed GitHub Actions workflow runs.", "POST /api/v1/agent/sessions/:sessionId/actions/runs/:runId/debug."),
            ("GitHub Actions", "Prepares write actions and requires human approval before mutation.", "POST /api/v1/agent/sessions/:sessionId/github/write; POST /api/v1/agent/github/approvals/:approvalId/decision."),
        ],
    )

    add_heading(doc, "4. Backend Route Map")
    add_para(doc, "The backend registers two route groups in src/server.ts:")
    add_bullets(
        doc,
        [
            "/api/v1/agent -> agent.routes.ts for sessions, analysis, chat, CI debugging, and GitHub write approvals.",
            "/api/v1/github -> github.routes.ts for GitHub authentication status and repository listing.",
            "/api/v1/health -> health check endpoint for API availability.",
        ],
    )

    add_heading(doc, "5. Session and Repository Context")
    add_para(doc, "The most important backend concept is repository context. After repository selection, the backend stores owner, repo, fullName, defaultBranch, privacy flag, URL, and related metadata on the session.")
    add_para(doc, "All repository-aware services follow the same pattern:")
    add_numbered(
        doc,
        [
            "Read sessionId from the route or request body.",
            "Fetch the session from MongoDB.",
            "Read the selected repository from the session.",
            "Block early if no repository is selected.",
            "Pass the repository context into the agent run.",
        ],
    )
    add_para(doc, "This prevents undefined repository coordinates, which was the root cause behind GitHub calls like owner/repo being passed as undefined.")

    add_heading(doc, "6. AI Agent Responsibilities")
    add_kv_table(
        doc,
        [
            ("GitArchitect chat agent", "Answers user questions using selected repository context and GitHub evidence."),
            ("Repository discovery agent", "Inspects repository metadata and structure before deeper analysis."),
            ("Specialist agents", "Architecture, security, technology, and testing specialists analyze focused areas."),
            ("Aggregator agent", "Combines specialist outputs into one repository intelligence report."),
            ("PR review agent", "Reads PR metadata, diff, checks, files, and repository context to produce score and recommendation."),
            ("Issue analyzer agent", "Reads issue body/discussion and repository files to derive requirements, plan, affected areas, and questions."),
            ("CI debugger agent", "Reads workflow run/log evidence and repository config to identify likely failure causes."),
            ("GitHub write agent", "Prepares mutation actions but pauses for human approval before execution."),
        ],
    )

    add_heading(doc, "7. GitHub Integration")
    add_para(doc, "GitHub is used in two ways:")
    add_bullets(
        doc,
        [
            "GitHub REST API loads authenticated user repositories and verifies repository access.",
            "GitHub MCP tools provide repository-aware evidence for agents, including files, issues, pull requests, checks, Actions runs, and logs.",
        ],
    )
    add_para(doc, "The backend should never let agents analyze a repository only from its name. Agents are instructed to inspect actual GitHub evidence before making repository-specific claims.")

    add_heading(doc, "8. Human-Approved GitHub Writes")
    add_para(doc, "Write operations are intentionally separated from analysis. GitArchitect can propose a mutation, but the user must approve or reject it.")
    add_numbered(
        doc,
        [
            "User enters an instruction on the GitHub Actions page.",
            "Frontend calls /sessions/:sessionId/github/write.",
            "Backend resolves repository context and runs the GitHub write agent.",
            "If a write tool is proposed, backend stores an approval request in MongoDB.",
            "Frontend renders an approval card with repository, tool name, payload, Approve, and Reject.",
            "User decision calls /github/approvals/:approvalId/decision.",
            "Backend resumes or cancels the paused write operation.",
        ],
    )
    add_callout(
        doc,
        "Bug fix note.",
        "The approval API must receive a real approvalId. If the frontend calls /approvals/undefined/decision, normalize the backend response and guard the button action so the UI only submits when approvalId exists.",
        fill=COLORS["warn"],
    )

    add_heading(doc, "9. Frontend Rendering Rules")
    add_para(doc, "Several backend responses are nested. The frontend must unwrap the correct object before rendering cards.")
    add_kv_table(
        doc,
        [
            ("PR Review", "Backend response shape: success -> data -> review. UI cards should read review.pullRequest, review.summary, review.riskLevel, review.reviewRecommendation, review.scores, review.findings, review.filesReviewed, review.limitations."),
            ("Issue Analyzer", "Backend response shape: success -> data -> analysis. UI cards should read analysis.issue, analysis.issueType, analysis.summary, analysis.problemStatement, analysis.requirements, analysis.implementationPlan, analysis.testingPlan, analysis.risks, analysis.questions."),
            ("GitHub Writes", "Backend may return approvalId when approval is required. UI should display the approval card only when approvalId is present."),
            ("Empty states", "Cards should show meaningful fallback text like 'No findings reported' instead of blank sections."),
        ],
    )

    add_heading(doc, "10. Page-by-Page Presentation Notes")
    add_kv_table(
        doc,
        [
            ("Dashboard", "Start here. Explain that the app is session-based, repository-aware, and GitHub-connected."),
            ("Repositories", "Show that selecting a repository is the setup step that powers every later page."),
            ("AI Chat", "Explain that chat is not generic. It answers using selected repository context."),
            ("Repository Intelligence", "Explain this as the broad health report: architecture, security, testing, maintainability, and technology stack."),
            ("PR Review", "Explain how the top cards summarize decision, risk, score, changed files, checks, and summary."),
            ("PR Review details", "Explain the evidence sections: positives, findings, files reviewed, testing assessment, checks, and limitations."),
            ("Issue Analyzer", "Explain how it turns issue text into implementation-ready details and flags missing requirements."),
            ("Issue details", "Show requirements, affected areas, implementation plan, testing plan, risks, questions, and related files."),
            ("CI Debugger", "Explain that it reads GitHub Actions run/log evidence to identify root cause and recommended fix."),
            ("GitHub Actions", "Explain that mutations are human-approved. The agent proposes, the user decides, then backend executes or rejects."),
        ],
    )

    add_heading(doc, "11. Important Fixes Already Applied")
    add_bullets(
        doc,
        [
            "Backend service calls must await session/repository helpers before using repository fields.",
            "Issue analysis frontend must read response.data.analysis, not the top-level response object.",
            "PR review frontend must read response.data.review and render nested fields.",
            "Approval flow must pass approval._id or equivalent approvalId to the frontend.",
            "Frontend approve/reject action must guard against undefined approvalId before calling the decision endpoint.",
        ],
    )

    add_heading(doc, "12. Data Flow Example: Issue Analyzer")
    add_numbered(
        doc,
        [
            "User enters issue number 10.",
            "Frontend calls POST /api/v1/agent/sessions/:sessionId/issues/10/analyze.",
            "Backend loads selected repository from MongoDB session.",
            "Issue analyzer agent fetches GitHub issue details and inspects relevant repository files.",
            "Backend returns success.data.analysis.",
            "Frontend renders issue number, type, readiness, confidence, summary, problem statement, requirements, plan, risks, questions, and related files.",
        ],
    )

    add_heading(doc, "13. Data Flow Example: PR Review")
    add_numbered(
        doc,
        [
            "User enters PR number 9.",
            "Frontend calls POST /api/v1/agent/sessions/:sessionId/pull-requests/9/review.",
            "Backend loads selected repository from session.",
            "PR review agent fetches PR metadata, diff, file changes, checks, and relevant repository context.",
            "Backend returns success.data.review plus overallScore and repository metadata.",
            "Frontend renders recommendation, risk, score, PR title, branches, files reviewed, positives, findings, checks, limitations, and testing assessment.",
        ],
    )

    add_heading(doc, "14. Demo Script")
    add_numbered(
        doc,
        [
            "Open Dashboard and show GitHub connected, session active, selected repository, and human-approved write status.",
            "Open Repositories and explain that selecting a repository writes context to the backend session.",
            "Open PR Review and run PR #9 to show recommendation, risk, score, checks, files, and limitations.",
            "Open Issue Analyzer and run issue #10 to show requirements, missing details, plan, risks, questions, and related files.",
            "Open GitHub Actions and explain the approval card. Emphasize that write operations require explicit human approval.",
            "Close by explaining the backend pattern: Angular sends sessionId, Express resolves repository from MongoDB, agents inspect GitHub evidence, frontend renders structured cards.",
        ],
    )

    add_heading(doc, "15. Risks and Operational Notes")
    add_bullets(
        doc,
        [
            "If no repository is selected, repository-specific workflows must block with a clear message.",
            "If GitHub token permissions are limited, checks or private repository data may be unavailable.",
            "If backend returns nested data and frontend reads the wrong level, cards appear blank even when the API response is valid.",
            "If approvalId is missing, approve/reject calls will hit /approvals/undefined/decision and fail.",
            "If agents cannot inspect GitHub evidence, responses should report limitations instead of inventing details.",
        ],
    )

    add_heading(doc, "16. Files and Modules to Mention")
    add_kv_table(
        doc,
        [
            ("Backend entry", "src/server.ts"),
            ("Agent routes", "src/routes/agent.routes.ts"),
            ("GitHub routes", "src/routes/github.routes.ts"),
            ("Session logic", "src/controllers/session.controller.ts and src/services/session.service.ts"),
            ("Repository analysis", "src/controllers/repository-analyzer.controller.ts and repository intelligence services/agents"),
            ("PR review", "src/controllers/pr-review.controller.ts and src/services/pr-review.service.ts"),
            ("Issue analyzer", "src/controllers/issue-analyzer.controller.ts and src/services/issue-analyzer.service.ts"),
            ("CI debugger", "src/controllers/ci-debugger.controller.ts and src/services/ci-debugger.service.ts"),
            ("GitHub writes", "src/controllers/github-write.controller.ts and src/services/github-write.service.ts"),
            ("Frontend APIs", "frontend core API services call the backend endpoints and unwrap nested response data."),
        ],
    )

    add_heading(doc, "17. Technology Stack")
    add_bullets(
        doc,
        [
            "Angular frontend for pages, forms, cards, and result rendering.",
            "Node.js and Express backend for REST APIs.",
            "TypeScript across frontend and backend.",
            "MongoDB for sessions and approval state.",
            "OpenAI Agents SDK for AI workflows.",
            "Zod for structured output/schema validation.",
            "GitHub REST/MCP integration for repository evidence.",
        ],
    )

    add_heading(doc, "18. Final Team Message")
    add_para(doc, "GitArchitect is built around one consistent architecture: select repository once, store it in the session, resolve it in backend services, let specialized agents inspect GitHub evidence, and render structured results in the UI. The same pattern powers PR Review, Issue Analyzer, CI Debugger, Repository Intelligence, Chat, and human-approved GitHub write actions.")

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("GitArchitect application flow document")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = COLORS["muted"]

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
